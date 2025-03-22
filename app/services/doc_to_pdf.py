import os
from docx import Document
from xhtml2pdf import pisa  
import pypandoc
import docx.enum.style
import subprocess
import re

def markdown_to_docx(markdown_content, filename='fact_check_report.docx'):
    # Create a new Document
    doc = Document()

    # Add styles for list bullet
    styles = doc.styles
    if 'List Bullet' not in styles:
        styles.add_style('List Bullet', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

    # Split the markdown content into lines
    lines = markdown_content.split('\n')

    # Iterate through each line and add it to the document
    for line in lines:
        # Trim whitespace
        line = line.strip()
        
        # Check for headers using regex
        header_match = re.match(r'^(#+)\s*(.+)', line)
        if header_match:
            header_level = len(header_match.group(1))  # Count the number of '#' to determine level
            header_text = header_match.group(2)
            print(f"Adding Header Level {header_level}: {header_text}")  # Debugging output
            doc.add_heading(header_text, level=header_level)
            continue  # Skip the rest of the loop for headers

        # Handle bold text
        if "**" in line:
            para = doc.add_paragraph()
            split_text = line.split("**")
            for i, text in enumerate(split_text):
                run = para.add_run(text)
                if (i % 2) != 0:
                    run.bold = True
            continue  # Skip the rest of the loop for bold text

        # Check for list items
        if line.startswith("- "):
            para = doc.add_paragraph(line[2:], style='List Bullet')
            continue  # Skip the rest of the loop for list items

        # Add normal text
        doc.add_paragraph(line)  # Normal text

    # Save the document
    try:
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error saving docx file: {e}")
        return None


def convert_docx_to_pdf(docx_filepath):
    try:
        pdf_filepath = docx_filepath.replace(".docx", ".pdf")
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', docx_filepath, '--outdir', os.path.dirname(pdf_filepath)],
                       check=True)

        return pdf_filepath
    except subprocess.CalledProcessError as e:
        print(f"Error converting docx to pdf: {e}")
        return None