import streamlit as st
from dotenv import load_dotenv
from langchain.agents import initialize_agent, AgentType
from langchain.tools import Tool
from langchain_groq import ChatGroq
from langchain.tools import DuckDuckGoSearchRun, TavilySearchResults
from langchain_community.utilities import GoogleSerperAPIWrapper
from pydantic import BaseModel, HttpUrl, Field
from langchain.output_parsers import PydanticOutputParser
from langchain.schema import SystemMessage
from langchain.document_loaders import UnstructuredURLLoader
from datetime import datetime
import textwrap
import google.auth
from google.oauth2 import service_account
from googleapiclient.discovery import build
import google.generativeai as genai
import os
import re
from docx import Document
from docx2pdf import convert 
from docx.shared import Inches
from xhtml2pdf import pisa  
import pypandoc
import docx.enum.style
from googleapiclient.http import MediaFileUpload
import traceback


load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")  

GROQ= os.getenv("GROQ_API_KEY")

# Configure the genai library with your API key
genai.configure(api_key=GOOGLE_API_KEY)
llm_groq= ChatGroq(model='llama3-70b-8192', api_key= GROQ)
llm_gemini = genai.GenerativeModel(model_name="gemini-2.0-flash")


# Search functions
def google_search_tool(query):
    return GoogleSerperAPIWrapper().run(query)

def duckduckgo_tool(query):
    return DuckDuckGoSearchRun().run(query)

def tavily_tool(query):
    return TavilySearchResults(search_depth="advanced", include_answer=True, include_images=True).invoke(query)

# Wrap tools in Tool class
tools = [
    Tool(name="Google Search", func=google_search_tool, description="Search Google for news"),
    Tool(name="DuckDuckGo Search", func=duckduckgo_tool, description="Search DuckDuckGo for news"),
    Tool(name="Tavily Search", func=tavily_tool, description="Search Tavily for related sources")
]

# Initialize the agent
agent = initialize_agent(
    tools=tools,
    llm=llm_groq,
    agent_type=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
    verbose=True
)

# Define structured output using Pydantic
class FactCheckURLs(BaseModel):
    claim: str = Field(description="The claim being fact-checked.")
    urls: list[HttpUrl] = Field(description="A list of URLs that contain relevant news articles.")

# Initialize output parser
parser = PydanticOutputParser(pydantic_object=FactCheckURLs)

# Function to generate the fact-check report
def generate_report(fact_check_result, urls):
    current_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    report_header = f"""
{"#" * 50}
Fact-Check Report
Date: {current_date}
{"#" * 50}
"""
    claim_section = f"""
## Claim:
{user_claim}
"""
    fact_check_section = f"""
## Fact-Check Findings:
{fact_check_result}
"""
    sources_section = f"""
## Sources:
The following sources were used to verify the claim:
"""
    for url in urls:
        sources_section += f"- {url}\n"

    full_report = report_header + claim_section + fact_check_section + sources_section
    return full_report


def markdown_to_docx(markdown_content,filename= 'fact_check_report.docx'):  
    # Create a new Document
    doc = Document()

    # Add styles for list bullet (Conditionally)
    styles = doc.styles
    if 'List Bullet' not in styles:
        styles.add_style('List Bullet', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

    # Split the markdown content into lines
    lines = markdown_content.split('\n')

    # Iterate through each line and add it to the document
    for line in lines:
        if line.startswith("# "):  # Header level 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  # Header level 2
            doc.add_heading(line[3:], level=2)
        elif "**" in line: # Bold text
            para = doc.add_paragraph()
            # Split line into list of strings
            split_text = line.split("**")
            # Enumerate the list
            for i, text in enumerate(split_text):
                # If even, text is not bold, if odd it is bold
                run = para.add_run(text)
                if (i%2) != 0:
                    run.bold = True

        elif line.startswith("- "):  # List item
            para = doc.add_paragraph(line[2:], style='List Bullet')  # style applies to paragraph, not run

        else:
            doc.add_paragraph(line)  # Normal text

    # Save the document
    try:
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error saving docx file: {e}")
        return None



# Convert DOCX to PDF
def convert_docx_to_pdf(docx_filepath, pdf_filepath):
    try:
        # Convert docx to html
        output = pypandoc.convert_file(docx_filepath, 'html', outputfile="output.html")
        assert output == ""

        # Convert html to pdf
        with open("output.html", "r") as f:
            html = f.read()

        # Convert the HTML to PDF
        with open(pdf_filepath, "wb") as f:
            pisa_status = pisa.CreatePDF(html, dest=f)

        # Remove temporary html report
        os.remove("output.html")

        return pdf_filepath
    except Exception as e:
        print(f"Error converting docx to pdf: {e}")
        return None

# Load Google Drive credentials

# Load Google Drive credentials
def get_google_services():
    creds = service_account.Credentials.from_service_account_file("/teamspace/studios/this_studio/fact-checker-450613-deb87d989f62.json", scopes=["https://www.googleapis.com/auth/documents", "https://www.googleapis.com/auth/drive.file"])
    docs_service = build("docs", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds) # Drive API
    return docs_service, drive_service

# Extract Folder ID from full Google Drive folder link
def extract_folder_id(drive_link):
    match = re.search(r"drive\.google\.com/drive/folders/([a-zA-Z0-9_-]+)", drive_link)
    return match.group(1) if match else None

def upload_docx_to_drive(docx_filepath, drive_link):
    """Uploads the DOCX file to Google Drive and returns a shareable link."""
    try:
        docs_service, drive_service = get_google_services()

        # Extract Folder ID
        folder_id = extract_folder_id(drive_link)
        if not folder_id:
            print("Error: Invalid Google Drive folder link.")
            return None

        # File metadata
        file_metadata = {
            "name": os.path.basename(docx_filepath),
            "parents": [folder_id]
        }

        # Media upload
        media = MediaFileUpload(docx_filepath, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Upload the file
        try:
            file = drive_service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        except Exception as upload_error:
            print(f"Error during file upload: {upload_error}")
            print(traceback.format_exc())  # Print the full traceback
            return None

        # Get the file ID and webViewLink
        file_id = file.get('id')
        file_url = file.get('webViewLink')

        return file_url

    except Exception as e:
        print(f"Error uploading DOCX to Google Drive: {e}")
        print(traceback.format_exc())  # Print the full traceback
        return None

# ✅ User provides the full Google Drive folder link
drive_folder_link = "https://drive.google.com/drive/folders/1kPc5b41KpJur_pTKLXGLlCB2IjD5-eC2?dmr=1&ec=wgc-drive-hero-goto"

# Streamlit interface
st.title("📰 Fact-Checking Tool")
st.write("Enter a claim to fact-check and get a detailed report.")

# User input for the claim
user_claim = st.text_input("Enter the claim to fact-check:",)
if st.button("Fact-Check"):
    with st.spinner("Analyzing the claim..."):
        try:
            # Step 1: Retrieve relevant news sources
            query = f"Find credible news sources related to this claim: {user_claim}.{parser.get_format_instructions()}"
            fact_check_data = agent.run(query)

            # Step 2: Parse response
            fact_check_data = parser.parse(fact_check_data)
            try:
                if not isinstance(fact_check_data, str):  # Check if it's not a string
                    fact_check_data = fact_check_data.model_dump_json()  # Convert to JSON if not a string
                fact_check_data = parser.parse(fact_check_data)  # Parse the data
            except Exception as e:
                print(f"Error parsing response: {e}")
                print(f"Raw response: {fact_check_data}")
                print(f"Expected format: {parser.get_format_instructions()}")
                raise e 
            # Step 3: Load full content of the URLs
            urls_to_check = fact_check_data.urls
            loader = UnstructuredURLLoader(urls=urls_to_check)
            documents = loader.load()
            news_texts = [doc.page_content for doc in documents]

            # Step 4: Fact-check using GPT-4
            fact_check_prompt = f"""
            Claim: "{user_claim}"

            Below are news articles related to the claim:

            {news_texts[:2000]}  # Limit to avoid token overflow

            Analyze the claim using these sources:
            - Identify which parts are **true, false, or uncertain**.
            - Highlight contradictions between sources.
            - Reference sources when making a judgment.
            """
            fact_check_report = llm_gemini.generate_content(fact_check_prompt)
            fact_check_report= fact_check_report.text
            # Log the raw response object in Streamlit
            #st.write("Raw Fact-Check Response:")
            #st.write(fact_check_report)

            # Step 5: Generate report
            formatted_report = generate_report(fact_check_report, fact_check_data.urls)
      
            # Display report
            st.subheader("Fact-Check Report")
            st.write(formatted_report)

            st.subheader("Download Report")

            # PDF Download Button
            docx_filename = "fact_check_report.docx" # You must specify
            docx_file = markdown_to_docx(formatted_report, docx_filename)

            pdf_filename = "fact_check_report.pdf"
            pdf_file = convert_docx_to_pdf(docx_file, pdf_filename) # pdf_file will have the filename as before.

            if pdf_file:
                try:
                    with open(pdf_file, "rb") as file: # Open the actual pdf file
                        pdf_bytes = file.read()
                    st.download_button(
                        label="📄 Download Report as PDF",
                        data=pdf_bytes, # This must be the pdf bytes, not the docx
                        file_name=pdf_filename,
                        mime="application/pdf"
                    )
                except Exception as e:
                    st.error(f"Error reading PDF file: {e}")
            else:
                st.error("Failed to convert DOCX to PDF.")

            google_doc_url = upload_docx_to_drive(docx_file, drive_folder_link)

            if google_doc_url:
             # Button to open the Google Doc
             st.write("Access as google doc:", google_doc_url)

        except Exception as e:
            st.error(f"An error occurred: {e}")
            st.write("Raw response from the agent:", formatted_report)
