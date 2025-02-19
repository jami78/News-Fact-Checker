import os
from docx import Document

from xhtml2pdf import pisa  
import pypandoc
import docx.enum.style

def markdown_to_docx(markdown_content,filename= 'fact_check_report.docx'):  
    # Create a new Document
    doc = Document()

    # Add styles for list bullet (Conditionally)
    styles = doc.styles
    if 'List Bullet' not in styles:
        styles.add_style('List Bullet', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)

    # Split the markdown content into lines
    lines = markdown_content.split('\n')

    # Iterate through each line and add it to the document
    for line in lines:
        if line.startswith("# "):  # Header level 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  # Header level 2
            doc.add_heading(line[3:], level=2)
        elif "**" in line: # Bold text
            para = doc.add_paragraph()
            # Split line into list of strings
            split_text = line.split("**")
            # Enumerate the list
            for i, text in enumerate(split_text):
                # If even, text is not bold, if odd it is bold
                run = para.add_run(text)
                if (i%2) != 0:
                    run.bold = True

        elif line.startswith("- "):  # List item
            para = doc.add_paragraph(line[2:], style='List Bullet')  # style applies to paragraph, not run

        else:
            doc.add_paragraph(line)  # Normal text

    # Save the document
    try:
        doc.save(filename)
        return filename
    except Exception as e:
        print(f"Error saving docx file: {e}")
        return None



# Convert DOCX to PDF
def convert_docx_to_pdf(docx_filepath, pdf_filepath):
    try:
        # Convert docx to html
        output = pypandoc.convert_file(docx_filepath, 'html', outputfile="output.html")
        assert output == ""

        # Convert html to pdf
        with open("output.html", "r") as f:
            html = f.read()

        # Convert the HTML to PDF
        with open(pdf_filepath, "wb") as f:
            pisa_status = pisa.CreatePDF(html, dest=f)

        # Remove temporary html report
        os.remove("output.html")

        return pdf_filepath
    except Exception as e:
        print(f"Error converting docx to pdf: {e}")
        return None

